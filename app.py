# ==========================
# 🔧 Standard Library Imports
# ==========================
import os
import re
from datetime import datetime, timedelta
from io import BytesIO

import requests
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from bs4 import BeautifulSoup
import base64
import requests
from math import ceil
# ==========================
# 🧩 Third-Party Imports
# ==========================
from flask import (Flask, Response, flash, redirect, render_template, request,
                   send_from_directory, session, url_for)
from flask_sqlalchemy import SQLAlchemy


app = Flask(__name__)

load_dotenv()  # .env dosyasını yükler

app.secret_key = os.getenv("SECRET_KEY")
ADMIN_ID = os.getenv("ADMIN_ID")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD")

# Veritabanı bağlantısı
#app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///blog.db"
app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv("DATABASE_URL")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)

with app.app_context():
    db.create_all()

# -------------------------------
# MODEL TANIMI
# -------------------------------
class Post(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(150), nullable=False)
    content = db.Column(db.Text, nullable=False)
    date_posted = db.Column(db.DateTime, default=datetime.utcnow)
    views = db.Column(db.Integer, default=0)

    def __repr__(self):
        return f"<Post {self.title}>"

# -------------------------------
# ADMIN PANELİ GİRİŞ
# -------------------------------
@app.route("/zytez-login", methods=["GET", "POST"])
def zytez_login():
    """Gizli admin giriş ekranı."""
    if request.method == "POST":
        user_id = request.form.get("user_id")
        password = request.form.get("password")
        app.permanent_session_lifetime = timedelta(minutes=30)
        session.permanent = True

        if user_id == ADMIN_ID and password == ADMIN_PASSWORD:
            session["is_admin"] = True
            flash("✅ Giriş başarılı!", "success")
            return redirect(url_for("zytez"))
        else:
            flash("❌ Hatalı ID veya şifre!", "error")
            return redirect(url_for("zytez_login"))

    return render_template("zytez_login.html")

# -------------------------------
# ADMIN PANELİ SAYFASI
# -------------------------------
@app.route("/zytez")
def zytez():
    """Admin paneli (sadece giriş yapan görebilir)."""
    if not session.get("is_admin"):
        flash("🔒 Lütfen giriş yapın!", "warning")
        return redirect(url_for("zytez_login"))
    return render_template("zytez.html")

# -------------------------------
# ADD POST
# -------------------------------

@app.route("/zytez/add", methods=["POST"])
def zytez_add_post():
    """Admin panelinden yeni blog ekleme"""
    if not session.get("is_admin"):
        flash("🔒 Giriş yapmanız gerekiyor!", "warning")
        return redirect(url_for("zytez_login"))

    title = request.form.get("title")
    content = request.form.get("content")

    if not title or not content:
        flash("⚠️ Başlık ve içerik boş olamaz!", "error")
        return redirect(url_for("zytez"))

    new_post = Post(title=title, content=content)
    db.session.add(new_post)
    db.session.commit()
    flash("✅ Yeni yazı eklendi!", "success")

    return redirect(url_for("zytez"))

# -------------------------------
# DELETE POST
# -------------------------------

@app.route("/zytez/delete", methods=["POST"])
def zytez_delete_post():
    """Başlığa göre post silme"""
    if not session.get("is_admin"):
        flash("🔒 Giriş yapmanız gerekiyor!", "warning")
        return redirect(url_for("zytez_login"))

    title = request.form.get("title")

    if not title:
        flash("⚠️ Başlık boş olamaz!", "error")
        return redirect(url_for("zytez"))

    post = Post.query.filter_by(title=title).first()

    if not post:
        flash("❌ Bu başlığa sahip bir yazı bulunamadı!", "error")
    else:
        db.session.delete(post)
        db.session.commit()
        flash(f"🗑️ '{title}' başlıklı yazı silindi.", "success")

    return redirect(url_for("zytez"))

# -------------------------------
# DİĞER ROUTELAR
# -------------------------------
@app.route("/")
def index():
    page = request.args.get('page', 1, type=int)
    per_page = 20
    posts = Post.query.order_by(Post.date_posted.desc()).paginate(page=page, per_page=per_page)
    return render_template('index.html', posts=posts)

@app.route("/post/<int:id>")
def post(id):
    post_data = Post.query.get_or_404(id)
    post_data.views += 1
    db.session.commit()
    return render_template("post.html", post=post_data)

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/assets/<path:filename>")
def custom_static(filename):
    return send_from_directory(os.path.join(app.root_path, 'assets'), filename)

@app.route("/robots.txt")
def robots():
    return send_from_directory(os.path.join(app.root_path, 'static'), 'robots.txt')

@app.route("/search")
def search():
    query = request.args.get("q", "").lower()
    results = Post.query.filter(Post.title.ilike(f"%{query}%")).all()
    return render_template("search.html", results=results, query=query)


@app.route('/backup')
def backup():
    posts = Post.query.order_by(Post.date_posted.desc()).all()
    doc = Document()
    doc.add_heading('📚 Blog Yedek Raporu', 0)

    for post in posts:
        doc.add_heading(post.title, level=1)
        doc.add_paragraph(f"Tarih: {post.date_posted.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph('')

        # HTML'i parse et
        soup = BeautifulSoup(post.content, 'html.parser')

        # Hem p hem img etiketlerini sırayla gez
        for element in soup.find_all(['p', 'img']):
            if element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)

            elif element.name == 'img':
                src = element.get('src')
                if not src:
                    continue

                # Base64 mü, URL mi kontrol et
                if src.startswith('data:image'):
                    try:
                        header, data = src.split(',', 1)
                        data = data.replace('\n', '')  # Temizlik
                        image_data = base64.b64decode(data)
                        doc.add_picture(BytesIO(image_data), width=Inches(4))
                    except Exception as e:
                        doc.add_paragraph(f"[Görsel çözülemedi: {e}]")

                elif src.startswith('http'):
                    try:
                        img_data = requests.get(src, timeout=10).content
                        doc.add_picture(BytesIO(img_data), width=Inches(4))
                    except Exception as e:
                        doc.add_paragraph(f"[Görsel indirilemedi: {src}]")

        doc.add_page_break()

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return Response(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=blog_backup.docx"}
    )

@app.route("/sitemap.xml", methods=["GET"])
def sitemap():
    """Tüm blog yazıları için otomatik site haritası üretir."""
    posts = Post.query.order_by(Post.date_posted.desc()).all()
    base_url = "https://ytez-abap-blog.onrender.com"

    # Son post tarihi (ya da bugünün tarihi)
    lastmod = posts[0].date_posted if posts else datetime.utcnow()

    xml = render_template("sitemap.xml", posts=posts, base_url=base_url, lastmod=lastmod)
    return Response(xml, mimetype="application/xml")

# -------------------------------
# ÇIKIŞ (Opsiyonel)
# -------------------------------
@app.route("/logout")
def logout():
    session.pop("is_admin", None)
    flash("👋 Oturum kapatıldı.", "info")
    return redirect(url_for("index"))

# -------------------------------
# MAIN
# -------------------------------
if __name__ == "__main__":
    app.run(debug=False)